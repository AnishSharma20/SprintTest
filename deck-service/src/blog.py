"""Blog draft generator — turns source material (uploaded files and/or picked scientific-study
summaries) into an on-brand, science-based Superba Krill blog draft in Markdown, in the style of
superbakrill.com/blog.

Recipe distilled from Aker BioMarine's own blogs (e.g. the weight-loss/muscle-loss article):
~1,100–1,900 words, B2B + informed-consumer audience, ~70% scientific authority / 30% marketing,
problem-hook intro → 4–6 H2 sections (problem → nutritional opportunity → how Superba works →
clinical evidence → broader benefits → science-supported solution) → whitepaper/contact CTA, with
named studies, real stats and numbered references. Claim fidelity is enforced (same rules as decks).
"""
from __future__ import annotations

import re

import anthropic

from . import config
from .planner import CLAIM_RULES


def strip_dashes(md: str) -> str:
    """Deterministic safety net for the no-dash brand rule: remove every dash from the article
    TEXT, while preserving Markdown structure (leading '- '/'* ' bullets and '---' rules).
    Em/en dashes become a comma; a hyphen between word characters becomes a space
    (so "Omega-3" -> "Omega 3", "double-blind" -> "double blind")."""
    out = []
    for line in md.splitlines():
        if line.strip() in ("---", "***", "___"):
            out.append(line)
            continue
        m = re.match(r"^(\s*[-*]\s+)", line)  # keep a leading list marker intact
        prefix, rest = (m.group(1), line[m.end():]) if m else ("", line)
        rest = re.sub(r"\s*[—–]\s*", ", ", rest)      # em/en dash -> comma
        rest = re.sub(r"(?<=\w)-(?=\w)", " ", rest)   # inter-word/number hyphen -> space
        out.append(prefix + rest)
    return "\n".join(out)

WORDS = {"kort": "700–950", "standard": "1100–1500", "detaljert": "1600–1950"}

TONE_GUIDANCE = {
    "salg":       "Lean benefit-first (~55% marketing / 45% science): lead each section with the reader/market benefit, use the studies as proof beneath it. Persuasive but never overclaiming.",
    "balansert":  "Balance benefit and evidence (~30% marketing / 70% science) — the default Superba blog voice: credible and authoritative, benefits always tied to the science.",
    "vitenskap":  "Evidence-led (~15% marketing / 85% science): thorough on study designs, numbers and mechanisms; measured, HCP-facing. Still readable, still ends on the practical takeaway.",
}


def build_system(length: str, tone: str, instructions: str = "") -> str:
    words = WORDS.get(length, WORDS["standard"])
    instr = ""
    if (instructions or "").strip():
        instr = ("\n\nUSER CONTEXT & INSTRUCTIONS (high-priority guidance on audience, angle, emphasis, "
                 "terminology, what to include/avoid — follow it unless it conflicts with the claim-fidelity "
                 "rules below):\n\"\"\"\n" + instructions.strip() + "\n\"\"\"\n")
    return f"""You are a science content writer for Aker BioMarine's Superba Krill. Write a publish-ready
BLOG DRAFT in Markdown, in the exact style of the superbakrill.com/blog articles, based ONLY on the
source material provided (scientific study summaries and/or documents).

AUDIENCE: supplement brand owners, formulators and health professionals, plus informed consumers.
Credible and authoritative, but accessible — explain the science, don't dumb it down.

TONE: {TONE_GUIDANCE.get(tone, TONE_GUIDANCE['balansert'])}

LENGTH: about {words} words.

STRUCTURE (Markdown):
- One `#` H1: a benefit- or problem-framed headline (like "How Krill Oil Can Help Address a Hidden Risk
  in Weight Loss: Muscle Loss"). Not a bland topic label.
- A 2–3 short-paragraph intro that opens on a problem, market trend or reader pain — then previews the answer.
- 4–6 `##` H2 sections. A proven arc: the problem/context → the nutritional opportunity → HOW Superba Krill
  works (phospholipid-bound EPA/DHA, choline, astaxanthin; superior absorption vs fish oil) → THE CLINICAL
  EVIDENCE (the studies in the source) → broader benefits → a short "a science-supported solution" close.
- Use **bold** for key phrases, bullet lists for benefits/mechanisms, and call out the strongest statistics.
- End with a short **call to action**: invite the reader to download the matching Superba whitepaper and/or
  contact Aker BioMarine.
- A final `## References` section: a numbered list of the studies actually cited (author, year, journal if given).

USING THE SCIENCE (critical):
- Ground the blog in the studies present in the source. Name the study type (e.g. "6-month randomized,
  double-blind, placebo-controlled trial"), sample sizes, and the real figures (%, effect sizes, p-values)
  AS STATED in the source. Attribute with author + year.
- {CLAIM_RULES}
- Do NOT invent studies, numbers, quotes or references. If the source is thin on a point, keep it general
  rather than fabricating specifics. Distinguish clearly between what a study showed and general mechanism.

LANGUAGE: if the user context below specifies an output language, write the ENTIRE blog in that language;
otherwise write in the same language as the source material. Keep brand names (Superba, Aker BioMarine) as-is.

TEXT STYLE (strict brand rule): do NOT use dash characters anywhere in the article TEXT. Never an em-dash,
an en-dash, or a hyphen between words. Rephrase to avoid them (write "evidence based", "double blind",
"Omega 3", "phospholipid bound", "12 week"); use commas, colons, parentheses or separate words instead.
Markdown list markers ("- ") and a "---" divider are structure, not text, and are fine.
{instr}
Output ONLY the Markdown blog draft, with no preamble, no explanation and no code fences."""


def _title(markdown: str, fallback: str) -> str:
    for line in markdown.splitlines():
        if line.startswith("# "):
            return line[2:].strip()
    return fallback


def generate_blog(client: anthropic.Anthropic, source_text: str, base_name: str, *,
                  length: str = "standard", tone: str = "balansert", instructions: str = "",
                  on_progress=None) -> dict:
    if on_progress:
        try:
            on_progress(10, "Reading the source & studies")
        except Exception:  # noqa: BLE001
            pass
    msg = client.messages.create(
        model=config.MODEL, max_tokens=6000, system=build_system(length, tone, instructions),
        messages=[{"role": "user", "content":
                   f"SOURCE MATERIAL:\n{source_text}\n\nWrite the Superba Krill blog draft now, in Markdown."}],
    )
    if on_progress:
        try:
            on_progress(90, "Finalizing the draft")
        except Exception:  # noqa: BLE001
            pass
    markdown = "".join(b.text for b in msg.content if getattr(b, "type", "") == "text").strip()
    # strip an accidental ```markdown fence if the model added one
    if markdown.startswith("```"):
        markdown = markdown.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
    markdown = strip_dashes(markdown)  # enforce the no-dash brand rule deterministically
    return {"markdown": markdown, "filename": f"{base_name}.md",
            "title": _title(markdown, base_name)}


# ---------------------------------------------------------------------------
# Markdown -> Word (.docx). The draft is generated and reviewed as Markdown, but the
# deliverable teams want is a Word document. Small, dependency-light converter built on
# python-docx (already a dependency): headings, paragraphs, bullet/numbered lists, and
# inline **bold** / *italic*. Good enough for a clean, editable Word draft.
# ---------------------------------------------------------------------------
def markdown_to_docx(markdown: str, title: str | None = None) -> bytes:
    import io
    import re

    import docx

    doc = docx.Document()

    def add_inline(paragraph, text: str) -> None:
        # [label](url) -> "label (url)"; then split out **bold** and *italic* runs.
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", lambda m: f"{m.group(1)} ({m.group(2)})", text)
        for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True
            elif part.startswith("*") and part.endswith("*"):
                paragraph.add_run(part[1:-1]).italic = True
            else:
                paragraph.add_run(part)

    for raw in markdown.splitlines():
        s = raw.strip()
        if not s:
            continue
        if s in ("---", "***", "___"):
            continue
        if s.startswith("#### "):
            doc.add_heading(s[5:].strip(), level=3)
        elif s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=2)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=1)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=0)  # Title style
        elif re.match(r"^[-*]\s+", s):
            add_inline(doc.add_paragraph(style="List Bullet"), re.sub(r"^[-*]\s+", "", s))
        elif re.match(r"^\d+\.\s+", s):
            add_inline(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s+", "", s))
        else:
            add_inline(doc.add_paragraph(), s)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
